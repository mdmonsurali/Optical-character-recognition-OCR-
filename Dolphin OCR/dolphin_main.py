import argparse
import glob
import os
import cv2
import torch
from PIL import Image
from transformers import AutoProcessor, VisionEncoderDecoderModel
import json

import uuid
import base64
from io import BytesIO

from docx import Document
from docx.shared import Inches
from bs4 import BeautifulSoup

import sys
sys.path.append(os.path.join(os.path.dirname(__file__), "Dolphin"))


from utils.utils import *


class DOLPHIN:
    def __init__(self, model_id_or_path):
        """Initialize the Hugging Face model
        
        Args:
            model_id_or_path: Path to local model or Hugging Face model ID
        """
        # Load model from local path or Hugging Face hub
        self.processor = AutoProcessor.from_pretrained(model_id_or_path)
        self.model = VisionEncoderDecoderModel.from_pretrained(model_id_or_path)
        self.model.eval()
        
        # Set device and precision
        self.device = "cuda" if torch.cuda.is_available() else "cpu"
        self.model.to(self.device)
        self.model = self.model.half()  # Always use half precision by default
        
        # set tokenizer
        self.tokenizer = self.processor.tokenizer
        
    def chat(self, prompt, image):
        """Process an image or batch of images with the given prompt(s)
        
        Args:
            prompt: Text prompt or list of prompts to guide the model
            image: PIL Image or list of PIL Images to process
            
        Returns:
            Generated text or list of texts from the model
        """
        # Check if we're dealing with a batch
        is_batch = isinstance(image, list)
        
        if not is_batch:
            # Single image, wrap it in a list for consistent processing
            images = [image]
            prompts = [prompt]
        else:
            # Batch of images
            images = image
            prompts = prompt if isinstance(prompt, list) else [prompt] * len(images)
        
        # Prepare image
        batch_inputs = self.processor(images, return_tensors="pt", padding=True)
        batch_pixel_values = batch_inputs.pixel_values.half().to(self.device)
        
        # Prepare prompt
        prompts = [f"<s>{p} <Answer/>" for p in prompts]
        batch_prompt_inputs = self.tokenizer(
            prompts,
            add_special_tokens=False,
            return_tensors="pt"
        )

        batch_prompt_ids = batch_prompt_inputs.input_ids.to(self.device)
        batch_attention_mask = batch_prompt_inputs.attention_mask.to(self.device)
        
        # Generate text
        outputs = self.model.generate(
            pixel_values=batch_pixel_values,
            decoder_input_ids=batch_prompt_ids,
            decoder_attention_mask=batch_attention_mask,
            min_length=1,
            max_length=4096,
            pad_token_id=self.tokenizer.pad_token_id,
            eos_token_id=self.tokenizer.eos_token_id,
            use_cache=True,
            bad_words_ids=[[self.tokenizer.unk_token_id]],
            return_dict_in_generate=True,
            do_sample=False,
            num_beams=1,
            repetition_penalty=1.1
        )
        
        # Process output
        sequences = self.tokenizer.batch_decode(outputs.sequences, skip_special_tokens=False)
        
        # Clean prompt text from output
        results = []
        for i, sequence in enumerate(sequences):
            cleaned = sequence.replace(prompts[i], "").replace("<pad>", "").replace("</s>", "").strip()
            results.append(cleaned)
            
        # Return a single result for single image input
        if not is_batch:
            return results[0]
        return results


def process_page(image_path, model, save_dir, max_batch_size=None):
    """Parse document images with two stages"""
    # Stage 1: Page-level layout and reading order parsing
    pil_image = Image.open(image_path).convert("RGB")
    layout_output = model.chat("Parse the reading order of this document.", pil_image)

    # Stage 2: Element-level content parsing
    padded_image, dims = prepare_image(pil_image)
    recognition_results = process_elements(layout_output, padded_image, dims, model, max_batch_size)

    # Save outputs
    json_path = save_outputs(recognition_results, image_path, save_dir)

    return json_path, recognition_results


def process_elements(layout_results, padded_image, dims, model, max_batch_size=None):
    """Parse all document elements with parallel decoding"""
    layout_results = parse_layout_string(layout_results)

    # Store text and table elements separately
    text_elements = []  # Text elements
    table_elements = []  # Table elements
    figure_results = []  # Image elements (no processing needed)
    previous_box = None
    reading_order = 0

    # Collect elements to process and group by type
    for bbox, label in layout_results:
        try:
            # Adjust coordinates
            x1, y1, x2, y2, orig_x1, orig_y1, orig_x2, orig_y2, previous_box = process_coordinates(
                bbox, padded_image, dims, previous_box
            )

            # Crop and parse element
            cropped = padded_image[y1:y2, x1:x2]
            if cropped.size > 0:
                if label == "fig":
                    # For figure regions, add empty text result immediately
                    figure_results.append(
                        {
                            "label": label,
                            "bbox": [orig_x1, orig_y1, orig_x2, orig_y2],
                            "text": "",
                            "reading_order": reading_order,
                        }
                    )
                else:
                    # Prepare element for parsing
                    pil_crop = Image.fromarray(cv2.cvtColor(cropped, cv2.COLOR_BGR2RGB))
                    element_info = {
                        "crop": pil_crop,
                        "label": label,
                        "bbox": [orig_x1, orig_y1, orig_x2, orig_y2],
                        "reading_order": reading_order,
                    }
                    
                    # Group by type
                    if label == "tab":
                        table_elements.append(element_info)
                    else:  # Text elements
                        text_elements.append(element_info)

            reading_order += 1

        except Exception as e:
            print(f"Error processing bbox with label {label}: {str(e)}")
            continue

    # Initialize results list
    recognition_results = figure_results.copy()
    
    # Process text elements (in batches)
    if text_elements:
        text_results = process_element_batch(text_elements, model, "Read text in the image.", max_batch_size)
        recognition_results.extend(text_results)
    
    # Process table elements (in batches)
    if table_elements:
        table_results = process_element_batch(table_elements, model, "Parse the table in the image.", max_batch_size)
        recognition_results.extend(table_results)

    # Sort elements by reading order
    recognition_results.sort(key=lambda x: x.get("reading_order", 0))

    return recognition_results


def process_element_batch(elements, model, prompt, max_batch_size=None):
    """Process elements of the same type in batches"""
    results = []
    
    # Determine batch size
    batch_size = len(elements)
    if max_batch_size is not None and max_batch_size > 0:
        batch_size = min(batch_size, max_batch_size)
    
    # Process in batches
    for i in range(0, len(elements), batch_size):
        batch_elements = elements[i:i+batch_size]
        crops_list = [elem["crop"] for elem in batch_elements]
        
        # Use the same prompt for all elements in the batch
        prompts_list = [prompt] * len(crops_list)
        
        # Batch inference
        batch_results = model.chat(prompts_list, crops_list)
        
        # Add results
        for j, result in enumerate(batch_results):
            elem = batch_elements[j]
            results.append({
                "label": elem["label"],
                "bbox": elem["bbox"],
                "text": result.strip(),
                "reading_order": elem["reading_order"],
            })
    
    return results


def main():
    parser = argparse.ArgumentParser(description="Document processing tool using DOLPHIN model")
    parser.add_argument("--model_path", default="./hf_model", help="Path to Hugging Face model")
    parser.add_argument("--input_path", type=str, default="./demo", help="Path to input image or directory of images")
    parser.add_argument(
        "--save_dir",
        type=str,
        default=None,
        help="Directory to save parsing results (default: same as input directory)",
    )
    parser.add_argument(
        "--max_batch_size",
        type=int,
        default=16,
        help="Maximum number of document elements to parse in a single batch (default: 16)",
    )
    args = parser.parse_args()

    # Load Model
    model = DOLPHIN(args.model_path)

    # Collect Document Images
    if os.path.isdir(args.input_path):
        image_files = []
        for ext in [".jpg", ".jpeg", ".png", ".JPG", ".JPEG", ".PNG"]:
            image_files.extend(glob.glob(os.path.join(args.input_path, f"*{ext}")))
        image_files = sorted(image_files)
    else:
        if not os.path.exists(args.input_path):
            raise FileNotFoundError(f"Input path {args.input_path} does not exist")
        image_files = [args.input_path]

    save_dir = args.save_dir or (
        args.input_path if os.path.isdir(args.input_path) else os.path.dirname(args.input_path)
    )
    setup_output_dirs(save_dir)

    total_samples = len(image_files)
    print(f"\nTotal samples to process: {total_samples}")

    # Process All Document Images
    for image_path in image_files:
        print(f"\nProcessing {image_path}")
        try:
            json_path, recognition_results = process_page(
                image_path=image_path,
                model=model,
                save_dir=save_dir,
                max_batch_size=args.max_batch_size,
            )

            print(f"Processing completed. Results saved to {save_dir}")

        except Exception as e:
            print(f"Error processing {image_path}: {str(e)}")
            continue


#Add figure to json and markdown

def process_figures_and_generate_markdown(image_path, results, markdown_text, save_dir="cropped_figures"):
    os.makedirs(save_dir, exist_ok=True)
    image = Image.open(image_path)

    updated_results = []
    markdown_parts = []
    
    for entry in sorted(results, key=lambda x: x["reading_order"]):
        if entry["label"] == "fig":
            # Generate unique filename
            unique_name = f"fig_{uuid.uuid4().hex[:8]}.jpg"
            image_path_out = os.path.join(save_dir, unique_name)

            # Crop and save
            bbox = entry["bbox"]
            cropped = image.crop((bbox[0], bbox[1], bbox[2], bbox[3]))
            cropped.save(image_path_out)

            # Replace 'text' with 'img'
            entry["img"] = image_path_out
            entry.pop("text", None)

            # Convert image to base64 for markdown
            buffered = BytesIO()
            cropped.save(buffered, format="JPEG")
            img_b64 = base64.b64encode(buffered.getvalue()).decode("utf-8")
            img_md = f"![figure](data:image/jpeg;base64,{img_b64})"

            markdown_parts.append(img_md)
        else:
            markdown_parts.append(entry["text"])
            updated_results.append(entry)

    # Recombine markdown with double line breaks
    updated_markdown = "\n\n".join(markdown_parts)

    # Remove 'reading_order' from each item in JSON
    cleaned_json = [{k: v for k, v in item.items() if k != "reading_order"} for item in results]
    
    return cleaned_json, updated_markdown


def json_to_docx(json_data, output_path="output.docx", image_max_width=5, base_path="."):
    """
    Convert layout-aware JSON into a DOCX file.
    Supports 'sec', 'para', 'cap', 'fig' (with 'img'), and 'tab' (HTML tables).
    """
    doc = Document()

    for item in json_data:
        label = item.get("label")

        if label == "sec":
            doc.add_heading(item["text"].strip(), level=1)

        elif label == "para":
            doc.add_paragraph(item["text"].strip())

        elif label == "cap":
            doc.add_paragraph(item["text"].strip(), style="Intense Quote")

        elif label == "fig":
            image_path = item.get("img")
            if image_path:
                # Resolve full path
                full_image_path = os.path.join(base_path, image_path)
                if os.path.exists(full_image_path):
                    try:
                        doc.add_picture(full_image_path, width=Inches(image_max_width))
                    except Exception as e:
                        doc.add_paragraph(f"[Error displaying image: {image_path}] {e}")
                else:
                    doc.add_paragraph(f"[Missing Image File: {image_path}]")
            else:
                doc.add_paragraph("[No image path provided]")

        elif label == "tab":
            html_table = item.get("text", "")
            soup = BeautifulSoup(html_table, "html.parser")
            rows = soup.find_all("tr")

            if rows:
                num_cols = len(rows[0].find_all(["td", "th"]))
                table = doc.add_table(rows=0, cols=num_cols)
                table.style = 'Table Grid'

                for row in rows:
                    cells = row.find_all(["td", "th"])
                    doc_row = table.add_row().cells
                    for i, cell in enumerate(cells):
                        doc_row[i].text = cell.get_text(strip=True)

        else:
            # Unknown label fallback
            text = item.get("text") or item.get("img")
            if text:
                doc.add_paragraph(f"{label.upper()}: {text.strip()}")

    # Save to path or BytesIO buffer
    if hasattr(output_path, "write"):
        doc.save(output_path)
    else:
        doc.save(output_path)




def run_dolphin(image_path=None, image=None, save_dir=None, max_batch_size=16):
    if image_path:
        pil_image = Image.open(image_path).convert("RGB")
    elif image:
        pil_image = image.convert("RGB")
    else:
        raise ValueError("Either image_path or image must be provided")

    tmp_path = "./temp_input.jpg"
    pil_image.save(tmp_path)

    model_path = os.path.join(os.path.dirname(__file__), "Dolphin", "hf_model")
    model = DOLPHIN(model_path)

    save_dir = save_dir or "."
    setup_output_dirs(save_dir)

    json_path, recognition_results = process_page(
        image_path=tmp_path,
        model=model,
        save_dir=save_dir,
        max_batch_size=max_batch_size,
    )

    text_parts = [item["text"].strip() for item in recognition_results if item["text"].strip()]
    plain_text_output = "\n\n".join(text_parts)

    # Enhance with figures and markdown
    updated_json, updated_markdown = process_figures_and_generate_markdown(
        image_path=tmp_path,
        results=recognition_results,
        markdown_text=plain_text_output
    )

    return updated_json, updated_markdown


