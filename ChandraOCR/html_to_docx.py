"""
html_to_docx.py
Converts Chandra OCR HTML output to a .docx file.
Handles: tables, bold, italic, headings, lists, paragraphs, br, code, hr, math
"""

from pathlib import Path
from bs4 import BeautifulSoup, NavigableString, Tag
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def html_to_docx(html_path: str, output_path: str = None) -> str:
    """
    Convert a Chandra OCR HTML file to .docx.
    Output filename matches input filename.

    Args:
        html_path: Path to .html or .md file containing HTML from Chandra
        output_path: Optional output path. Defaults to same dir, .docx extension.

    Returns:
        Path to the created .docx file
    """
    html_path = Path(html_path)
    if output_path is None:
        output_path = html_path.with_suffix(".docx")
    else:
        output_path = Path(output_path)

    html_content = html_path.read_text(encoding="utf-8")
    return _convert(html_content, output_path)


def html_string_to_docx(html_content: str, output_path: str) -> str:
    """
    Convert HTML string directly to .docx.

    Args:
        html_content: HTML string (Chandra OCR output)
        output_path: Output .docx path

    Returns:
        Path to the created .docx file
    """
    return _convert(html_content, Path(output_path))


def _convert(html_content: str, output_path: Path) -> str:
    doc = Document()

    # Default font + margins
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    soup = BeautifulSoup(html_content, "html.parser")

    # If ocr_layout: process each top-level div block in order
    top_divs = soup.find_all("div", attrs={"data-label": True})
    if top_divs:
        for div in top_divs:
            label = div.get("data-label", "Text")
            _process_block(doc, div, label)
    else:
        # Plain HTML (ocr prompt) — process body or root
        body = soup.body if soup.body else soup
        _process_node(doc, body)

    doc.save(str(output_path))
    print(f"Saved DOCX: {output_path}")
    return str(output_path)


def _process_block(doc, div, label: str):
    """Process a single layout block div based on its label."""
    label = label.strip()

    if label in ("Page-Header", "Page-Footer"):
        return  # skip headers/footers

    if label == "Table":
        table_tag = div.find("table")
        if table_tag:
            _add_table(doc, table_tag)
        return

    if label in ("Section-Header", "h1", "h2", "h3", "h4", "h5"):
        heading_map = {
            "Section-Header": "Heading 2",
            "h1": "Heading 1", "h2": "Heading 2",
            "h3": "Heading 3", "h4": "Heading 4",
        }
        style = heading_map.get(label, "Heading 2")
        p = doc.add_paragraph(style=style)
        _fill_paragraph(p, div)
        return

    if label in ("List-Group",):
        _process_list(doc, div)
        return

    if label in ("Code-Block",):
        p = doc.add_paragraph()
        run = p.add_run(div.get_text())
        run.font.name = "Courier New"
        run.font.size = Pt(9)
        p.paragraph_format.left_indent = Inches(0.5)
        return

    if label == "Equation-Block":
        p = doc.add_paragraph()
        run = p.add_run(div.get_text())
        run.italic = True
        run.font.name = "Cambria Math"
        return

    # Default: Text, Caption, Footnote, Complex-Block, Figure, Form etc.
    _process_node(doc, div)


def _process_node(doc, node):
    """Recursively process HTML nodes into docx."""
    for child in node.children:
        if isinstance(child, NavigableString):
            text = child.strip()
            if text:
                doc.add_paragraph(text)
            continue

        tag = child.name.lower() if child.name else ""

        if tag in ("h1", "h2", "h3", "h4", "h5", "h6"):
            level = int(tag[1])
            style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3",
                         4: "Heading 4", 5: "Heading 5", 6: "Heading 6"}
            p = doc.add_paragraph(style=style_map.get(level, "Heading 3"))
            _fill_paragraph(p, child)

        elif tag == "p":
            p = doc.add_paragraph()
            _fill_paragraph(p, child)

        elif tag == "table":
            _add_table(doc, child)

        elif tag in ("ul", "ol"):
            _process_list(doc, child)

        elif tag == "hr":
            _add_hr(doc)

        elif tag == "br":
            doc.add_paragraph()

        elif tag in ("div", "section", "article", "body"):
            _process_node(doc, child)

        elif tag in ("pre", "code"):
            p = doc.add_paragraph()
            run = p.add_run(child.get_text())
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            p.paragraph_format.left_indent = Inches(0.5)

        elif tag in ("b", "strong", "i", "em", "u", "span", "a",
                     "small", "big", "sub", "sup", "del", "math"):
            # Inline elements at block level — wrap in paragraph
            p = doc.add_paragraph()
            _fill_paragraph(p, child)

        else:
            # Unknown tag — recurse
            _process_node(doc, child)


def _fill_paragraph(paragraph, node):
    """Fill a paragraph with inline content from an HTML node."""
    for child in node.children:
        if isinstance(child, NavigableString):
            text = str(child)
            if text:
                paragraph.add_run(text)
            continue

        tag = child.name.lower() if child.name else ""

        if tag in ("b", "strong"):
            run = paragraph.add_run(child.get_text())
            run.bold = True

        elif tag in ("i", "em"):
            run = paragraph.add_run(child.get_text())
            run.italic = True

        elif tag == "u":
            run = paragraph.add_run(child.get_text())
            run.underline = True

        elif tag == "del":
            run = paragraph.add_run(child.get_text())
            run.font.strike_through = True

        elif tag == "sup":
            run = paragraph.add_run(child.get_text())
            run.font.superscript = True

        elif tag == "sub":
            run = paragraph.add_run(child.get_text())
            run.font.subscript = True

        elif tag == "code":
            run = paragraph.add_run(child.get_text())
            run.font.name = "Courier New"
            run.font.size = Pt(9)

        elif tag == "math":
            run = paragraph.add_run(child.get_text())
            run.italic = True
            run.font.name = "Cambria Math"

        elif tag == "a":
            run = paragraph.add_run(child.get_text())
            run.font.color.rgb = RGBColor(0x00, 0x56, 0xB2)
            run.underline = True

        elif tag == "br":
            paragraph.add_run("\n")

        elif tag in ("span", "small", "big", "label"):
            _fill_paragraph(paragraph, child)

        elif tag == "img":
            alt = child.get("alt", "")
            if alt:
                run = paragraph.add_run(f"[Image: {alt}]")
                run.italic = True
                run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        elif tag == "input":
            checked = child.get("checked") is not None
            input_type = child.get("type", "text")
            if input_type in ("checkbox",):
                paragraph.add_run("☑ " if checked else "☐ ")
            elif input_type == "radio":
                paragraph.add_run("● " if checked else "○ ")
            else:
                val = child.get("value", "")
                paragraph.add_run(f"[{val}]")

        else:
            _fill_paragraph(paragraph, child)


def _process_list(doc, list_node):
    """Process ul/ol and nested lists."""
    tag = list_node.name.lower() if list_node.name else "ul"
    style = "List Bullet" if tag == "ul" else "List Number"

    for child in list_node.children:
        if isinstance(child, NavigableString):
            continue
        if child.name and child.name.lower() == "li":
            # Check for nested list
            nested = child.find(["ul", "ol"])
            if nested:
                # Add the text part first
                text_parts = []
                for c in child.children:
                    if isinstance(c, NavigableString):
                        text_parts.append(str(c))
                    elif c.name and c.name.lower() not in ("ul", "ol"):
                        text_parts.append(c.get_text())
                p = doc.add_paragraph(style=style)
                _fill_paragraph(p, child)
                # Process nested list
                _process_list(doc, nested)
            else:
                p = doc.add_paragraph(style=style)
                _fill_paragraph(p, child)


def _add_table(doc, table_tag):
    """Convert HTML table to docx table."""
    rows = table_tag.find_all("tr")
    if not rows:
        return

    # Count max columns
    max_cols = 0
    for row in rows:
        cols = len(row.find_all(["td", "th"]))
        if cols > max_cols:
            max_cols = cols

    if max_cols == 0:
        return

    table = doc.add_table(rows=len(rows), cols=max_cols)
    table.style = "Table Grid"

    for r_idx, row in enumerate(rows):
        cells = row.find_all(["td", "th"])
        for c_idx, cell in enumerate(cells):
            if c_idx >= max_cols:
                break
            docx_cell = table.cell(r_idx, c_idx)
            docx_cell.text = ""
            p = docx_cell.paragraphs[0]
            p.clear()
            _fill_paragraph(p, cell)

            # Bold header cells
            is_header = cell.name.lower() == "th" or r_idx == 0
            if is_header:
                for run in p.runs:
                    run.bold = True
                    run.font.size = Pt(9)
            else:
                for run in p.runs:
                    run.font.size = Pt(9)

    doc.add_paragraph()  # spacing after table


def _add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ── CLI ────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    import sys
    if len(sys.argv) < 2:
        print("Usage: python html_to_docx.py input.html [output.docx]")
        sys.exit(1)
    inp = sys.argv[1]
    out = sys.argv[2] if len(sys.argv) > 2 else None
    html_to_docx(inp, out)